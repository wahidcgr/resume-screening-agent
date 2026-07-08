"""
parser.py
Extracts raw text from resume files in PDF, DOCX, or plain-text format.
"""

from pathlib import Path

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


def extract_text(file_path: str) -> str:
    """Read a resume file and return its raw text content.

    Supports .pdf, .docx, and .txt. Raises ValueError for anything else.
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext == ".txt":
        return _read_txt(path)
    elif ext == ".docx":
        return _read_docx(path)
    elif ext == ".pdf":
        return _read_pdf(path)
    else:
        raise ValueError(
            f"Unsupported file type '{ext}' for {path.name}. "
            f"Supported types: {sorted(SUPPORTED_EXTENSIONS)}"
        )


def _read_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_docx(path: Path) -> str:
    import docx  # python-docx

    doc = docx.Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also pull text out of any tables (some resumes use table layouts).
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)

    return "\n".join(parts)


def _read_pdf(path: Path) -> str:
    import pdfplumber

    parts = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            if text.strip():
                parts.append(text)
    return "\n".join(parts)


def list_resume_files(folder: str) -> list:
    """Return sorted list of resume file paths with supported extensions."""
    folder_path = Path(folder)
    files = [
        str(p) for p in sorted(folder_path.iterdir())
        if p.suffix.lower() in SUPPORTED_EXTENSIONS and p.is_file()
    ]
    return files
